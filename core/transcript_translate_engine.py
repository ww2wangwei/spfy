import importlib
import json
import re
import urllib.error
import urllib.request
from pathlib import Path
from typing import Dict, List

from docx import Document

translator_engine = importlib.import_module('video_tool.core.translator_engine')


class TranscriptTranslateEngine:
    """逐字稿批量翻译为 Word 表格"""

    def load_transcript(self, file_path: str) -> str:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.docx':
            doc = Document(file_path)
            return '\n'.join(p.text.strip() for p in doc.paragraphs if p.text.strip())

        if suffix == '.doc':
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(path.absolute()))
                content = doc.Content.Text
                doc.Close(False)
                return content.strip()
            finally:
                word.Quit()

        for encoding in ('utf-8', 'utf-8-sig', 'gbk', 'gb18030'):
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                continue

        raise Exception(f"无法识别文本文件编码: {file_path}")

    def split_transcript(self, text: str) -> List[str]:
        text = text.replace('\r', '\n')
        paragraphs = [p.strip() for p in re.split(r'\n+', text) if p.strip()]
        segments = []

        for paragraph in paragraphs:
            paragraph = re.sub(r'\s+', ' ', paragraph).strip()
            if not paragraph:
                continue
            segments.extend(self._split_paragraph(paragraph))

        return [seg for seg in segments if seg.strip()]

    def _split_paragraph(self, paragraph: str) -> List[str]:
        parts = []
        current = ''
        i = 0

        while i < len(paragraph):
            current += paragraph[i]
            should_break = False

            if paragraph[i] in '。！？；;':
                should_break = True
            elif paragraph[i] == '.' and self._is_sentence_period(paragraph, i):
                should_break = True
            elif len(current) >= 80 and paragraph[i] in '，, ':
                should_break = True

            if should_break:
                parts.append(current.strip())
                current = ''

            i += 1

        if current.strip():
            parts.append(current.strip())

        refined = []
        for part in parts:
            refined.extend(self._split_long_step_text(part))
        return refined

    def _is_sentence_period(self, text: str, index: int) -> bool:
        before = text[index - 1:index]
        after = text[index + 1:index + 2]
        if before.isalpha() and after.isalpha():
            return False
        return True

    def _split_long_step_text(self, text: str) -> List[str]:
        if len(text) <= 110:
            return [text]

        markers = [
            r'(?=第\s*\d+\s*步)',
            r'(?=第[一二三四五六七八九十]+步)',
            r'(?=首先)',
            r'(?=然后)',
            r'(?=接着)',
            r'(?=点击)',
        ]
        pattern = '|'.join(markers)
        pieces = [p.strip() for p in re.split(pattern, text) if p.strip()]
        if len(pieces) <= 1:
            return self._split_by_length(text)

        result = []
        buffer = ''
        for piece in pieces:
            if not buffer:
                buffer = piece
            elif len(buffer) + len(piece) <= 110:
                buffer += piece
            else:
                result.append(buffer.strip())
                buffer = piece
        if buffer:
            result.append(buffer.strip())
        return result

    def _split_by_length(self, text: str, max_len: int = 110) -> List[str]:
        result = []
        current = text
        while len(current) > max_len:
            split_at = max(
                current.rfind('，', 0, max_len),
                current.rfind(',', 0, max_len),
                current.rfind(' ', 0, max_len)
            )
            if split_at < max_len * 0.4:
                split_at = max_len
            result.append(current[:split_at + 1].strip())
            current = current[split_at + 1:].strip()
        if current:
            result.append(current)
        return result

    def translate_files(self, file_paths: List[str], progress_callback=None,
                        log_callback=None) -> List[str]:
        outputs = []
        total_files = len(file_paths)

        for file_index, file_path in enumerate(file_paths, 1):
            if log_callback:
                log_callback(f"开始翻译逐字稿: {file_path}")

            text = self.load_transcript(file_path)
            segments = self.split_transcript(text)

            if log_callback:
                log_callback(f"拆分为 {len(segments)} 个句段")

            rows = []
            for i, source in enumerate(segments, 1):
                if progress_callback:
                    base = int((file_index - 1) / total_files * 100)
                    span = int(100 / total_files)
                    progress = base + int(i / max(len(segments), 1) * span * 0.85)
                    progress_callback(progress, 100, f"翻译逐字稿 {file_index}/{total_files}: {i}/{len(segments)}")

                target = self._translate_segment(source, log_callback=log_callback)
                rows.append({'index': i, 'source': source, 'target': target})

            if progress_callback:
                progress_callback(
                    int(file_index / total_files * 100),
                    100,
                    f"保存翻译文档 {file_index}/{total_files}..."
                )

            output_path = self._save_docx(file_path, rows)
            outputs.append(output_path)

            if log_callback:
                log_callback(f"✓ 已输出: {output_path}")

        if progress_callback:
            progress_callback(100, 100, "逐字稿翻译完成")

        return outputs

    def _translate_segment(self, text: str, log_callback=None) -> str:
        model_key = translator_engine.get_translation_model()
        if model_key == "Google翻译(备用)":
            raise Exception("逐字稿翻译需要使用设置里的大模型，请不要选择 Google翻译(备用)")

        minimax_model = translator_engine.TRANSLATION_MODELS.get(model_key, "abab6.5s-chat")
        prompt = f"""Translate the Chinese transcript segment into natural English for a training course handout.

Rules:
1. Preserve technical terms, model names, software menu names, signal names, and code-like identifiers.
2. Keep the English clear and instructional.
3. Do not add explanations.
4. Output only the English translation.

Chinese:
{text}

English:"""

        payload = {
            "model": minimax_model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.2,
            "max_tokens": 800
        }

        try:
            req = urllib.request.Request(
                f"{translator_engine.MINIMAX_BASE_URL}/chat/completions",
                data=json.dumps(payload).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {translator_engine.MINIMAX_API_KEY}'
                },
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=120) as response:
                result = json.loads(response.read().decode('utf-8'))
                translated = result['choices'][0]['message']['content'].strip()
                translated = re.sub(r'^(English:)\s*', '', translated, flags=re.IGNORECASE).strip()
                return translated
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8') if e.fp else ''
            raise Exception(f"{model_key} HTTP {e.code}: {error_body[:200]}")
        except Exception as e:
            raise Exception(f"逐字稿翻译失败: {str(e)}")

    def _save_docx(self, source_path: str, rows: List[Dict]) -> str:
        source = Path(source_path)
        output_path = str(source.with_name(f"{source.stem}-翻译.docx"))

        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = '序号'
        header_cells[1].text = '原文'
        header_cells[2].text = '翻译文'

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row['index'])
            cells[1].text = row['source']
            cells[2].text = row['target']

        doc.save(output_path)
        return output_path


transcript_translate_engine = TranscriptTranslateEngine()
